#!/usr/bin/env python3
"""Genere un .docx stylé par livrable (Times New Roman, interligne 1,15).

Style sobre : une seule couleur d'accent (bleu ardoise), titres hierarchises,
filets fins, page de couverture mise en page, en-tete + pied de page numerote,
sommaire automatique.

Pipeline :
  1. Gabarit de reference (.docx) : polices, interligne, couleurs des titres.
  2. pandoc rend le CORPS (sections, hors couverture et sommaire manuel) + TOC auto.
  3. python-docx ajoute la couverture, l'en-tete, le pied de page et les filets.
"""
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

BASE = Path(__file__).resolve().parent
FONT = "Times New Roman"
LINE_SPACING = 1.15

# Palette sobre : un seul accent + nuances de gris
ACCENT = RGBColor(0x1F, 0x38, 0x64)   # bleu ardoise (titres 1)
ACCENT2 = RGBColor(0x2E, 0x50, 0x80)  # bleu plus clair (titres 2)
GREY = RGBColor(0x59, 0x59, 0x59)     # gris (titres 3, pied de page)
HEX_ACCENT = "1F3864"
HEX_GREY = "B0B0B0"

DELIVERABLES = {
    "Dossier_de_Projet_Perinea.docx": {
        "folder": "dossier-de-projet",
        "label": "TITRE PROFESSIONNEL — CONCEPTEUR DÉVELOPPEUR D'APPLICATIONS",
        "kind": "DOSSIER DE PROJET",
        "subtitle": "Application mobile de rééducation périnéale connectée",
        "footer": "Dossier de Projet — Périnéa",
    },
    "Dossier_Professionnel_Perinea.docx": {
        "folder": "dossier-professionnel",
        "label": "TITRE PROFESSIONNEL — CONCEPTEUR DÉVELOPPEUR D'APPLICATIONS",
        "kind": "DOSSIER PROFESSIONNEL",
        "subtitle": "Présentation des compétences — Projet Périnéa",
        "footer": "Dossier Professionnel — Périnéa",
    },
}

# Fichiers exclus du corps (gerés autrement) : couverture + sommaire manuel
EXCLUDE = ("00-page-de-garde", "02-sommaire")


# --------------------------------------------------------------------------- #
# Helpers bas niveau
# --------------------------------------------------------------------------- #
def _set_style_font(style, size=12, color=None, bold=None, italic=None):
    style.font.name = FONT
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)


def _par_border(paragraph, edge="bottom", sz=6, color=HEX_ACCENT, space=2):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        pPr.append(pbdr)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    pbdr.append(el)


def _add_field(paragraph, instr):
    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = instr
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(i); run._r.append(e)
    return run


def _style_run(run, size=11, color=None, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)


# --------------------------------------------------------------------------- #
# Gabarit de reference
# --------------------------------------------------------------------------- #
def build_reference():
    doc = Document()

    normal = doc.styles["Normal"]
    _set_style_font(normal, size=12)
    pf = normal.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    headings = {
        "Heading 1": (15, ACCENT, True, False),
        "Heading 2": (13, ACCENT2, True, False),
        "Heading 3": (11.5, GREY, True, True),
        "Heading 4": (11, GREY, False, True),
        "Title": (26, ACCENT, True, False),
        "Subtitle": (13, GREY, False, True),
    }
    for name, (size, color, bold, italic) in headings.items():
        if name in doc.styles:
            st = doc.styles[name]
            _set_style_font(st, size=size, color=color, bold=bold, italic=italic)
            st.paragraph_format.line_spacing = LINE_SPACING
            st.paragraph_format.space_before = Pt(10)
            st.paragraph_format.space_after = Pt(4)
            st.paragraph_format.keep_with_next = True

    ref = BASE / "_reference.docx"
    doc.save(ref)
    return ref


# --------------------------------------------------------------------------- #
# Corps via pandoc
# --------------------------------------------------------------------------- #
def concat_markdown(folder: Path) -> Path:
    files = [f for f in sorted(folder.glob("*.md"))
             if not f.stem.endswith(EXCLUDE)]
    pagebreak = ('\n\n```{=openxml}\n<w:p><w:r><w:br w:type="page"/>'
                 '</w:r></w:p>\n```\n\n')
    parts = []
    for i, f in enumerate(files):
        if i > 0:
            parts.append(pagebreak)
        parts.append(f.read_text(encoding="utf-8"))
    merged = folder.parent / f"_merged_{folder.name}.md"
    merged.write_text("\n\n".join(parts), encoding="utf-8")
    return merged


def render_body(merged: Path, out: Path, ref: Path):
    cmd = [
        "pandoc", str(merged),
        "-f", "gfm+raw_attribute",
        "-o", str(out),
        "--reference-doc", str(ref),
        "--toc", "--toc-depth=2",
        "--metadata", "toc-title=Sommaire",
    ]
    subprocess.run(cmd, check=True)


# --------------------------------------------------------------------------- #
# Couverture + en-tete/pied
# --------------------------------------------------------------------------- #
def _new_par():
    p = OxmlElement("w:p")
    return p


def build_cover(doc, meta):
    """Construit les paragraphes de couverture et les insere en tete du document."""
    created = []

    def para(text="", size=12, color=None, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6,
             border=False):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = LINE_SPACING
        if text:
            r = p.add_run(text)
            _style_run(r, size=size, color=color, bold=bold, italic=italic)
        if border:
            _par_border(p, edge="bottom", sz=8, color=HEX_ACCENT, space=4)
        created.append(p._p)
        return p

    para(space_before=36, space_after=2)
    para(meta["label"], size=10.5, color=GREY, bold=True, space_after=2)
    para(border=True, space_after=18)  # filet d'accent

    para(meta["kind"], size=22, color=ACCENT, bold=True, space_after=24)

    para("PROJET", size=12, color=GREY, bold=False, space_after=0)
    para("PÉRINÉA", size=40, color=ACCENT, bold=True, space_after=4)
    para(meta["subtitle"], size=13, color=GREY, italic=True, space_after=28)

    para(border=True, space_after=16)

    # Tableau d'informations (sobre)
    rows = [
        ("Candidat", "[À COMPLÉTER : Prénom NOM]"),
        ("Titre visé", "Concepteur Développeur d'Applications (CDA) — Niveau 6"),
        ("Session", "[À COMPLÉTER : mois / année]"),
        ("Centre de formation", "[À COMPLÉTER]"),
        ("Tuteur / Formateur", "[À COMPLÉTER]"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r, (k, v) in enumerate(rows):
        c0, c1 = table.rows[r].cells
        c0.text = ""; c1.text = ""
        rp = c0.paragraphs[0].add_run(k)
        _style_run(rp, size=11, color=ACCENT, bold=True)
        vp = c1.paragraphs[0].add_run(v)
        _style_run(vp, size=11)
        c0.paragraphs[0].paragraph_format.space_after = Pt(2)
        c1.paragraphs[0].paragraph_format.space_after = Pt(2)
    created.append(table._tbl)

    foot = para("Document imprimé et relié en deux exemplaires pour le jury.",
                size=10, color=GREY, italic=True, space_before=30, space_after=0)

    # Saut de page final
    pb = doc.add_paragraph()
    br_run = pb.add_run()
    br = OxmlElement("w:br"); br.set(qn("w:type"), "page")
    br_run._r.append(br)
    created.append(pb._p)

    # Deplace tout en tete du corps (dans l'ordre)
    body = doc.element.body
    for el in reversed(created):
        body.remove(el)
        body.insert(0, el)


def add_header_footer(doc, meta):
    section = doc.sections[0]
    section.different_first_page_header_footer = True  # couverture sans en-tete/pied

    # En-tete (pages de contenu)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("PÉRINÉA")
    _style_run(hr, size=9, color=GREY, bold=True)
    _par_border(hp, edge="bottom", sz=4, color=HEX_GREY, space=2)

    # Pied de page (pages de contenu) : titre a gauche, "Page X / Y" a droite
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = ""
    _par_border(fp, edge="top", sz=4, color=HEX_GREY, space=2)
    tabs = fp.paragraph_format.tab_stops
    # tabulation droite en bord de page (~16,5 cm)
    from docx.shared import Cm
    tabs.add_tab_stop(Cm(16.5))
    left = fp.add_run(meta["footer"] + "\t")
    _style_run(left, size=9, color=GREY)
    _add_field(fp, "PAGE")
    mid = fp.add_run(" / ")
    _style_run(mid, size=9, color=GREY)
    _add_field(fp, "NUMPAGES")
    for r in fp.runs:
        _style_run(r, size=9, color=GREY)


def style_heading_rules(doc):
    """Ajoute un filet fin sous les titres de niveau 1."""
    for p in doc.paragraphs:
        if p.style.name == "Heading 1":
            _par_border(p, edge="bottom", sz=6, color=HEX_ACCENT, space=2)


# --------------------------------------------------------------------------- #
def main():
    ref = build_reference()
    for out_name, meta in DELIVERABLES.items():
        folder = BASE / meta["folder"]
        merged = concat_markdown(folder)
        out = BASE / out_name
        render_body(merged, out, ref)
        merged.unlink()

        doc = Document(out)
        style_heading_rules(doc)
        build_cover(doc, meta)
        add_header_footer(doc, meta)
        doc.save(out)
        print(f"OK -> {out.name}")
    ref.unlink()


if __name__ == "__main__":
    main()
